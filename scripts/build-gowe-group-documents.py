from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
PROJECT_DIR = ROOT_DIR / "demo-data" / "gowe-group"
FILE_DIR = PROJECT_DIR / "file"
SOURCE_DATA_PATH = PROJECT_DIR / "source-data.json"


PRESET = {
    "name": "compact_reference_guide",
    "font": "Calibri",
    "body_size": 11,
    "body_after": 6,
    "line_spacing": 1.25,
    "h1_size": 16,
    "h1_color": "2E74B5",
    "h2_size": 13,
    "h2_color": "2E74B5",
    "h3_size": 12,
    "h3_color": "1F4D78",
    "table_header_fill": "E8EEF5",
}


def load_data() -> dict:
    with SOURCE_DATA_PATH.open("r", encoding="utf-8") as file:
        return json.load(file)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(value))
    run.font.name = PRESET["font"]
    run.font.size = Pt(9 if len(str(value)) > 120 else 10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table) -> None:
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin in ("top", "bottom", "start", "end"):
                node = tc_mar.find(qn(f"w:{margin}"))
                if node is None:
                    node = OxmlElement(f"w:{margin}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "120" if margin in ("start", "end") else "80")
                node.set(qn("w:type"), "dxa")


def apply_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = PRESET["font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["line_spacing"]

    for style_name, size, color in (
        ("Heading 1", PRESET["h1_size"], PRESET["h1_color"]),
        ("Heading 2", PRESET["h2_size"], PRESET["h2_color"]),
        ("Heading 3", PRESET["h3_size"], PRESET["h3_color"]),
    ):
        style = document.styles[style_name]
        style.font.name = PRESET["font"]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(7 if style_name != "Heading 3" else 5)


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.font.name = PRESET["font"]
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = PRESET["font"]
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor.from_string("606266")


def add_source_note(document: Document, urls: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Sources: " + "; ".join(urls))
    run.font.name = PRESET["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("606266")


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = PRESET["line_spacing"]
        paragraph.add_run(item)


def add_label_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = False
    table.rows[0].cells[0].width = Inches(1.875)
    table.rows[0].cells[1].width = Inches(4.625)
    set_table_width(table)
    for index, (label, value) in enumerate(rows):
        row = table.rows[0] if index == 0 else table.add_row()
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
        set_cell_shading(row.cells[0], PRESET["table_header_fill"])


def add_matrix_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    set_table_width(table)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], PRESET["table_header_fill"])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value)


def write_profile_doc(data: dict) -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(doc, "GOWE Group Profile", "Company profile source for Content Studio demo workflows")
    doc.add_heading("Company Overview", level=1)
    doc.add_paragraph(
        "GOWE Group is a construction formwork and scaffolding solution supplier integrating R&D, production, sales, leasing, construction, and operation across formwork, scaffolding, steel structures, bridge and tunnel equipment, and related building materials."
    )
    add_label_table(
        doc,
        [
            ("Company", "GOWE Group"),
            ("Website", "https://www.gowe-group.com/"),
            ("Industry", "Construction formwork, scaffolding, steel structure, bridge and tunnel equipment"),
            ("Positioning", "One-stop formwork and scaffolding partner for global construction teams."),
            ("Contact", "HW-GOWE@gowe-group.com; +86 18222117137; WhatsApp +8618222117137"),
            (
                "Address",
                "Unit 901-911, North Tower, Citic Poly Plaza, No.46 Tianhong Road, Lecong Town, Shunde, Foshan, Guangdong, China",
            ),
        ],
    )
    doc.add_heading("Core Capabilities", level=1)
    add_bullets(
        doc,
        [
            "Formwork and scaffolding R&D, production, sales, rental, construction support, and operation.",
            "Bridge and tunnel equipment, steel structure systems, protection platforms, and cantilever formwork.",
            "Engineering drawing support, material calculation, technical guidance, and localized overseas service.",
            "Project experience across residential, public, industrial, infrastructure, special building, and shipbuilding scenarios.",
        ],
    )
    doc.add_heading("Trust Signals", level=1)
    add_bullets(
        doc,
        [
            "30+ years of formwork and scaffolding experience.",
            "10,000+ construction projects and 800+ construction partners cited by the company.",
            "Certifications cited by GOWE include ISO 9001, ISO 14001, OHSAS 18001, EN 1090-2, and EN 12811.",
            "Overseas service network includes Thailand, Singapore, Malaysia, and broader localized branches.",
        ],
    )
    add_source_note(doc, ["https://www.gowe-group.com/", "https://www.gowe-group.com/about-us/", "https://www.gowe-group.com/contact/"])
    output = FILE_DIR / "GOWE Group Profile.docx"
    doc.save(output)
    return output


def write_product_doc(data: dict) -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(doc, "GOWE Product Portfolio", "Product portfolio source for catalog, product page, and distributor content")
    products = data["products"]
    categories = {}
    for product in products:
        categories.setdefault(product["category"], []).append(product)

    doc.add_heading("Portfolio Snapshot", level=1)
    add_matrix_table(
        doc,
        ["Category", "Products"],
        [[category, ", ".join(item["title"] for item in items)] for category, items in categories.items()],
    )
    for category, items in categories.items():
        doc.add_heading(category, level=1)
        for product in items:
            doc.add_heading(product["title"], level=2)
            doc.add_paragraph(product["summary"])
            add_source_note(doc, [product["sourceUrl"]])
    output = FILE_DIR / "GOWE Product Portfolio.docx"
    doc.save(output)
    return output


def write_solution_doc(data: dict) -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(doc, "GOWE Formwork and Scaffolding Solutions", "Lifecycle solution source for AI planning workflows")
    doc.add_heading("Service Capabilities", level=1)
    for service in data["services"]:
        doc.add_heading(service["title"], level=2)
        doc.add_paragraph(service["summary"])
        add_source_note(doc, [service["sourceUrl"]])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Lifecycle Solutions", level=1)
    add_matrix_table(
        doc,
        ["Solution", "Target Scenario", "Related Service"],
        [[item["title"], item["targetScenario"], item["relatedService"]] for item in data["solutions"]],
    )
    for solution in data["solutions"]:
        doc.add_heading(solution["title"], level=2)
        doc.add_paragraph(solution["summary"])
        add_source_note(doc, [solution["sourceUrl"]])

    output = FILE_DIR / "GOWE Formwork and Scaffolding Solutions.docx"
    doc.save(output)
    return output


def write_proof_faq_doc(data: dict) -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(doc, "GOWE Project Proof and FAQ", "Project proof and FAQ source for trust-building content")
    doc.add_heading("Project Proof", level=1)
    add_matrix_table(
        doc,
        ["Case", "Industry", "Summary"],
        [[item["title"], item["industry"], item["summary"]] for item in data["cases"]],
    )
    for case in data["cases"]:
        add_source_note(doc, [case["sourceUrl"]])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Frequently Asked Questions", level=1)
    for faq in data["faqs"]:
        doc.add_heading(faq["title"], level=2)
        doc.add_paragraph(faq["summary"])
        add_source_note(doc, [faq["sourceUrl"]])

    output = FILE_DIR / "GOWE Project Proof and FAQ.docx"
    doc.save(output)
    return output


def structural_check(paths: list[Path]) -> None:
    for path in paths:
        if not path.exists() or path.stat().st_size == 0:
            raise RuntimeError(f"Document was not written: {path}")
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "word/document.xml" not in names:
                raise RuntimeError(f"DOCX missing word/document.xml: {path}")


def main() -> None:
    FILE_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    paths = [
        write_profile_doc(data),
        write_product_doc(data),
        write_solution_doc(data),
        write_proof_faq_doc(data),
    ]
    structural_check(paths)
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
